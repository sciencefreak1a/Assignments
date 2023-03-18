#1. In what modes should the PdfFileReader() and PdfFileWriter() File objects will be opened?

#When using PdfFileReader() and PdfFileWriter() from the PyPDF2 library in Python, the file objects should be opened in binary mode by adding a 'b' to the file mode.
#This is because PDF files are binary files, and they need to be read and written in binary mode to ensure proper handling of the file's contents.
#Therefore, when opening a PDF file for reading using PdfFileReader(), the file should be opened in read-binary mode as follows:
import PyPDF2
from PyPDF2 import PdfFileReader

pdf_file = open(r"C:\Users\ajit1\Downloads\Tutorial_EDIT.pdf",'rb')
pdf_reader = PyPDF2.PdfReader(pdf_file)

#Similarly, when opening a file for writing using PdfFileWriter(), the file should be opened in write-binary mode:

import PyPDF2
writer = PyPDF2.PdfWriter()
with open(r"C:\Users\ajit1\Downloads\Tutorial_EDIT.pdf",'wb') as f:
    writer.write(f)

#2. From a PdfFileReader object, how do you get a Page object for page 5?

#To get a Page object for a specific page from a PdfFileReader object, you can use the len(reader.pages) method and pass the page number as an argument (0-indexed).
# Here's an example:

import PyPDF2
with open("C:\\Users\\ajit1\\Downloads\\Lecture_11.pdf",'rb') as p:
    reader = PyPDF2.PdfReader(p)
    #page = reader.pages[4] # get Page object for page 5 (0-indexed)
    num_pages = len(reader.pages)
    page = reader.pages[4]


#3. What PdfFileReader variable stores the number of pages in the PDF document?
#The PdfFileReader variable that stores the number of pages in the PDF document is numPages, which is an attribute of the PdfFileReader class.
import PyPDF2
pdf_file1 = open("C:\\Users\\ajit1\\Downloads\\Lecture_11.pdf",'rb')
pdf_reader = PyPDF2.PdfReader(pdf_file1)
num_pages1 = len(reader.pages)
print("The number of pages is : ", num_pages1)
pdf_file1.close()

#4. If a PdfFileReader object’s PDF is encrypted with the password swordfish, what must you do before you can obtain Page objects from it?
import PyPDF2

pdf_file = open("C:\\Users\\ajit1\\Downloads\\Lecture_11.pdf", 'rb')
pdf_reader = PyPDF2.PdfReader(pdf_file)
if pdf_reader.is_encrypted:
    pdf_reader.decrypt('swordfish')
page = reader.pages[4]

#5. What methods do you use to rotate a page?
from PyPDF2 import PdfReader, PdfWriter

# Open the PDF file
with open("C:\\Users\\ajit1\\Downloads\\Lecture_11.pdf", 'rb') as file:
    reader = PdfReader(file)

    # Get the first page
    page = reader.pages[0]

    # Rotate the page 90 degrees clockwise
    page.rotate(90)

    # Write the updated page to a new PDF file
    writer = PdfWriter()
    writer.add_page(page)

    with open('updated.pdf', 'wb') as output:
        writer.write(output)


#6. What is the difference between a Run object and a Paragraph object?
#In Python's python-docx library, a Paragraph object represents a single paragraph in a Word document, ' \
#while a Run object represents a contiguous run of text within a paragraph that has the same style and formatting.

#Here are some key differences between the two:

#A Paragraph object contains one or more Run objects.
#A Paragraph object has its own style and formatting attributes, such as alignment, indentation, and spacing, while a Run object inherits its style and formatting from the parent Paragraph object.
#A Paragraph object can contain other elements such as tables, lists, and images, while a Run object only contains text.
#A Paragraph object can be modified as a whole, while a Run object can only be modified within the context of its parent Paragraph object.
#In summary, a Paragraph object is used to represent and modify a whole paragraph in a Word document, while a Run object is used to represent and modify a contiguous run of text within a paragraph.

#7. How do you obtain a list of Paragraph objects for a Document object that’s stored in a variable named doc?
import docx
doc = docx.Document(r"C:\Users\ajit1\Downloads\file-sample_100kB.docx")
paragraph = doc.paragraphs

for paragraph in doc.paragraphs:
    print(paragraph.text)


#8. What type of object has bold, underline, italic, strike, and outline variables?
#The Run object in python-docx has the following variables:

#bold
#underline
#italic
#strike
#outline
#These variables are used to set or get the formatting properties of text in a Run object.

#9. What is the difference between False, True, and None for the bold variable?
#The bold variable can take three possible values: True, False, or None.

#True: If the bold variable is set to True, then the text of the associated Run object is formatted to be bold.
#False: If the bold variable is set to False, then the text of the associated Run object is not formatted to be bold.
#None: If the bold variable is set to None, then the text of the associated Run object inherits the bold formatting from the style applied to the containing Paragraph object.
# If no style is applied, the text is not bold.
#In summary, True means explicitly bold, False means explicitly not bold, and None means inherit bold formatting from style or not bold if no style is applied.

#10. How do you create a Document object for a new Word document?
import docx
doc = docx.Document()

doc.add_heading('About Myself',level = 0)
doc.add_paragraph("I am a 9 year old body. I go to school everyday. I play,eat,study and sleep")


doc.save('about_myself1.docx')

#11. How do you add a paragraph with the text "Hello, there!"; to a Document object stored in a variable named doc?
doc = docx.Document()

doc.add_paragraph("Hello, there!")

doc.save("question11.docx")

#12. What integers represent the levels of headings available in Word documents?
#In Word documents, the integers 1 to 9 represent the levels of headings available.
# Heading 1 is the highest level and Heading 9 is the lowest level.

